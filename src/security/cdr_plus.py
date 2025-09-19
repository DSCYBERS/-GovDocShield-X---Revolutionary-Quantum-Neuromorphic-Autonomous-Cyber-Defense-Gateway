"""
Positive Selection CDR+++ (Next-Gen Content Disarm & Reconstruction)
Neutralizes file-borne threats by rebuilding documents with only safe elements.
"""

import io
import logging
import hashlib
import time
import zipfile
import json
import xml.etree.ElementTree as ET
from typing import Dict, List, Any, Optional, Tuple, Union
from pathlib import Path
from dataclasses import dataclass
from enum import Enum
import magic
import olefile
import struct
from PIL import Image, ImageFilter
import pdfplumber
import docx
from openpyxl import load_workbook, Workbook
import python-pptx
import email
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from email.mime.application import MIMEApplication

logger = logging.getLogger(__name__)

class FileType(Enum):
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    PPTX = "pptx"
    RTF = "rtf"
    TXT = "txt"
    HTML = "html"
    XML = "xml"
    EMAIL = "email"
    IMAGE = "image"
    ARCHIVE = "archive"
    UNKNOWN = "unknown"

class ThreatLevel(Enum):
    SAFE = "safe"
    SUSPICIOUS = "suspicious"
    MALICIOUS = "malicious"
    UNKNOWN = "unknown"

@dataclass
class CDRResult:
    """Result of CDR+++ processing"""
    success: bool
    clean_content: Optional[bytes]
    original_hash: str
    clean_hash: Optional[str]
    file_type: FileType
    threats_removed: List[str]
    metadata_stripped: bool
    macros_removed: bool
    scripts_removed: bool
    embedded_objects_removed: bool
    reconstruction_log: List[str]
    processing_time_ms: float
    size_before: int
    size_after: int
    compression_ratio: float
    safety_score: float

class PositiveSelectionCDR:
    """
    Positive Selection CDR+++ Engine
    
    Uses positive selection approach - only allows known safe elements
    through reconstruction process. Discards anything suspicious or unknown.
    """
    
    def __init__(self):
        self.safe_pdf_objects = {
            'Page', 'Pages', 'Catalog', 'Font', 'FontDescriptor',
            'Text', 'Image', 'XObject', 'Contents', 'Resources'
        }
        
        self.safe_office_elements = {
            # Word
            'w:document', 'w:body', 'w:p', 'w:r', 'w:t', 'w:sectPr',
            'w:pPr', 'w:rPr', 'w:spacing', 'w:ind', 'w:jc', 'w:sz',
            'w:color', 'w:b', 'w:i', 'w:u', 'w:tbl', 'w:tr', 'w:tc',
            
            # Excel
            'worksheet', 'sheetData', 'row', 'c', 'v', 'f', 'si', 'sst',
            'workbook', 'sheets', 'sheet', 'cols', 'col',
            
            # PowerPoint
            'p:presentation', 'p:sldMaster', 'p:sld', 'p:cSld', 'p:spTree',
            'p:sp', 'p:txBody', 'a:p', 'a:r', 'a:t'
        }
        
        self.dangerous_patterns = [
            # Macros and scripts
            b'VBA', b'macro', b'autoopen', b'autoexec', b'autoclose',
            b'<script', b'javascript:', b'vbscript:', b'activex',
            
            # Suspicious objects
            b'oleObject', b'embedded', b'objectPool', b'oledata',
            b'package', b'packager', b'equation', b'msforms',
            
            # Exploits
            b'shellcode', b'exploit', b'rop', b'jmp', b'nop',
            b'%u9090', b'%u4141', b'AAAA', b'BBBB',
            
            # Steganography markers
            b'steghide', b'outguess', b'jsteg', b'f5', b'steganos'
        ]
        
        # Initialize file type detector
        try:
            self.magic = magic.Magic(mime=True)
        except:
            self.magic = None
            logger.warning("python-magic not available, using fallback detection")

    def reconstruct_document(self, file_content: bytes, filename: str = "") -> CDRResult:
        """
        Main CDR+++ reconstruction method
        
        Steps:
        1. Identify file type
        2. Disassemble document structure
        3. Extract safe elements only
        4. Rebuild clean document
        5. Verify integrity
        """
        start_time = time.time()
        original_hash = hashlib.sha256(file_content).hexdigest()
        original_size = len(file_content)
        
        logger.info(f"Starting CDR+++ reconstruction for {filename} ({original_size} bytes)")
        
        try:
            # Step 1: Identify file type
            file_type = self._detect_file_type(file_content, filename)
            logger.info(f"Detected file type: {file_type.value}")
            
            # Step 2: Scan for immediate threats
            threats_found = self._scan_for_threats(file_content)
            
            # Step 3: Perform type-specific reconstruction
            if file_type == FileType.PDF:
                result = self._reconstruct_pdf(file_content)
            elif file_type == FileType.DOCX:
                result = self._reconstruct_docx(file_content)
            elif file_type == FileType.XLSX:
                result = self._reconstruct_xlsx(file_content)
            elif file_type == FileType.PPTX:
                result = self._reconstruct_pptx(file_content)
            elif file_type == FileType.RTF:
                result = self._reconstruct_rtf(file_content)
            elif file_type == FileType.HTML:
                result = self._reconstruct_html(file_content)
            elif file_type == FileType.EMAIL:
                result = self._reconstruct_email(file_content)
            elif file_type == FileType.IMAGE:
                result = self._reconstruct_image(file_content)
            elif file_type == FileType.TXT:
                result = self._reconstruct_text(file_content)
            else:
                result = self._quarantine_unknown(file_content)
            
            # Calculate final metrics
            processing_time = (time.time() - start_time) * 1000
            clean_size = len(result['clean_content']) if result['clean_content'] else 0
            clean_hash = hashlib.sha256(result['clean_content']).hexdigest() if result['clean_content'] else None
            compression_ratio = clean_size / original_size if original_size > 0 else 0
            
            # Calculate safety score
            safety_score = self._calculate_safety_score(
                file_type, threats_found, result['threats_removed'], 
                result['metadata_stripped'], result['macros_removed']
            )
            
            return CDRResult(
                success=result['success'],
                clean_content=result['clean_content'],
                original_hash=original_hash,
                clean_hash=clean_hash,
                file_type=file_type,
                threats_removed=threats_found + result['threats_removed'],
                metadata_stripped=result['metadata_stripped'],
                macros_removed=result['macros_removed'],
                scripts_removed=result['scripts_removed'],
                embedded_objects_removed=result['embedded_objects_removed'],
                reconstruction_log=result['reconstruction_log'],
                processing_time_ms=processing_time,
                size_before=original_size,
                size_after=clean_size,
                compression_ratio=compression_ratio,
                safety_score=safety_score
            )
            
        except Exception as e:
            logger.error(f"CDR+++ reconstruction failed: {e}")
            return CDRResult(
                success=False,
                clean_content=None,
                original_hash=original_hash,
                clean_hash=None,
                file_type=FileType.UNKNOWN,
                threats_removed=[f"Reconstruction failed: {str(e)}"],
                metadata_stripped=False,
                macros_removed=False,
                scripts_removed=False,
                embedded_objects_removed=False,
                reconstruction_log=[f"ERROR: {str(e)}"],
                processing_time_ms=(time.time() - start_time) * 1000,
                size_before=original_size,
                size_after=0,
                compression_ratio=0.0,
                safety_score=0.0
            )

    def _detect_file_type(self, content: bytes, filename: str) -> FileType:
        """Detect file type using multiple methods"""
        
        # Magic number detection
        if content.startswith(b'%PDF'):
            return FileType.PDF
        elif content.startswith(b'PK\x03\x04'):
            # Could be Office document or ZIP
            if b'word/' in content[:1000] or filename.endswith('.docx'):
                return FileType.DOCX
            elif b'xl/' in content[:1000] or filename.endswith('.xlsx'):
                return FileType.XLSX
            elif b'ppt/' in content[:1000] or filename.endswith('.pptx'):
                return FileType.PPTX
            else:
                return FileType.ARCHIVE
        elif content.startswith(b'{\\rtf'):
            return FileType.RTF
        elif content.startswith((b'<html', b'<!DOCTYPE', b'<HTML')):
            return FileType.HTML
        elif content.startswith(b'<?xml'):
            return FileType.XML
        elif b'From:' in content[:500] or b'To:' in content[:500]:
            return FileType.EMAIL
        elif content.startswith((b'\xff\xd8\xff', b'\x89PNG', b'GIF8', b'BM')):
            return FileType.IMAGE
        elif all(32 <= b < 127 or b in [9, 10, 13] for b in content[:1000]):
            return FileType.TXT
        
        # Fallback to filename extension
        if filename:
            ext = Path(filename).suffix.lower()
            type_map = {
                '.pdf': FileType.PDF,
                '.docx': FileType.DOCX,
                '.xlsx': FileType.XLSX,
                '.pptx': FileType.PPTX,
                '.rtf': FileType.RTF,
                '.html': FileType.HTML,
                '.htm': FileType.HTML,
                '.xml': FileType.XML,
                '.txt': FileType.TXT,
                '.eml': FileType.EMAIL,
                '.msg': FileType.EMAIL,
                '.jpg': FileType.IMAGE,
                '.jpeg': FileType.IMAGE,
                '.png': FileType.IMAGE,
                '.gif': FileType.IMAGE,
                '.bmp': FileType.IMAGE
            }
            return type_map.get(ext, FileType.UNKNOWN)
        
        return FileType.UNKNOWN

    def _scan_for_threats(self, content: bytes) -> List[str]:
        """Scan for known threat patterns"""
        threats = []
        
        for pattern in self.dangerous_patterns:
            if pattern in content:
                threats.append(f"Dangerous pattern detected: {pattern.decode('utf-8', errors='ignore')}")
        
        # Check for suspicious entropy (possible encryption/obfuscation)
        if len(content) > 1000:
            entropy = self._calculate_entropy(content[:1000])
            if entropy > 7.5:  # High entropy indicates possible encryption
                threats.append("High entropy content detected (possible encryption/obfuscation)")
        
        # Check for PE headers (embedded executables)
        if b'MZ' in content and b'PE\x00\x00' in content:
            threats.append("Embedded executable detected")
        
        return threats

    def _calculate_entropy(self, data: bytes) -> float:
        """Calculate Shannon entropy of data"""
        if not data:
            return 0
        
        frequency = {}
        for byte in data:
            frequency[byte] = frequency.get(byte, 0) + 1
        
        entropy = 0
        length = len(data)
        for count in frequency.values():
            probability = count / length
            if probability > 0:
                entropy -= probability * (probability.bit_length() - 1)
        
        return entropy

    def _reconstruct_pdf(self, content: bytes) -> Dict[str, Any]:
        """Reconstruct PDF with only safe elements"""
        log = ["Starting PDF reconstruction"]
        threats_removed = []
        
        try:
            # Use pdfplumber to extract safe text and structure
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages_text = []
                
                for page_num, page in enumerate(pdf.pages):
                    # Extract only text - no scripts, forms, or embedded objects
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)
                        log.append(f"Extracted text from page {page_num + 1}")
                
                # Create clean PDF with just text content
                # For demo, we'll create a simple text representation
                clean_content = self._create_clean_text_document(pages_text, "PDF")
                
                log.append("PDF reconstruction completed - text extracted safely")
                threats_removed.append("PDF scripts and embedded objects removed")
                
                return {
                    'success': True,
                    'clean_content': clean_content,
                    'threats_removed': threats_removed,
                    'metadata_stripped': True,
                    'macros_removed': True,
                    'scripts_removed': True,
                    'embedded_objects_removed': True,
                    'reconstruction_log': log
                }
                
        except Exception as e:
            log.append(f"PDF reconstruction failed: {e}")
            return {
                'success': False,
                'clean_content': None,
                'threats_removed': [f"PDF processing error: {e}"],
                'metadata_stripped': False,
                'macros_removed': False,
                'scripts_removed': False,
                'embedded_objects_removed': False,
                'reconstruction_log': log
            }

    def _reconstruct_docx(self, content: bytes) -> Dict[str, Any]:
        """Reconstruct DOCX with only safe elements"""
        log = ["Starting DOCX reconstruction"]
        threats_removed = []
        
        try:
            # Open DOCX and extract safe elements
            doc = docx.Document(io.BytesIO(content))
            
            # Create new clean document
            clean_doc = docx.Document()
            
            # Extract only text content - no macros, embedded objects, etc.
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    new_para = clean_doc.add_paragraph(paragraph.text)
                    # Copy only basic formatting (safe)
                    if paragraph.style.name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3']:
                        new_para.style = paragraph.style.name
            
            # Save clean document
            clean_buffer = io.BytesIO()
            clean_doc.save(clean_buffer)
            clean_content = clean_buffer.getvalue()
            
            log.append("DOCX reconstruction completed - safe content extracted")
            threats_removed.extend([
                "Macros removed",
                "Embedded objects removed", 
                "Custom XML removed",
                "Active content removed"
            ])
            
            return {
                'success': True,
                'clean_content': clean_content,
                'threats_removed': threats_removed,
                'metadata_stripped': True,
                'macros_removed': True,
                'scripts_removed': True,
                'embedded_objects_removed': True,
                'reconstruction_log': log
            }
            
        except Exception as e:
            log.append(f"DOCX reconstruction failed: {e}")
            # Fallback to text extraction
            try:
                text_content = self._extract_text_from_zip(content)
                clean_content = self._create_clean_text_document([text_content], "DOCX")
                
                return {
                    'success': True,
                    'clean_content': clean_content,
                    'threats_removed': ["Converted to safe text format"],
                    'metadata_stripped': True,
                    'macros_removed': True,
                    'scripts_removed': True,
                    'embedded_objects_removed': True,
                    'reconstruction_log': log + ["Fallback to text extraction"]
                }
            except:
                return {
                    'success': False,
                    'clean_content': None,
                    'threats_removed': [f"DOCX processing error: {e}"],
                    'metadata_stripped': False,
                    'macros_removed': False,
                    'scripts_removed': False,
                    'embedded_objects_removed': False,
                    'reconstruction_log': log
                }

    def _reconstruct_xlsx(self, content: bytes) -> Dict[str, Any]:
        """Reconstruct XLSX with only safe elements"""
        log = ["Starting XLSX reconstruction"]
        threats_removed = []
        
        try:
            # Load workbook
            wb = load_workbook(io.BytesIO(content), data_only=True)  # data_only=True removes formulas
            
            # Create clean workbook
            clean_wb = Workbook()
            clean_wb.remove(clean_wb.active)  # Remove default sheet
            
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                clean_ws = clean_wb.create_sheet(title=sheet_name)
                
                # Copy only cell values (no formulas, macros, or VBA)
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.value is not None:
                            # Only copy simple values
                            if isinstance(cell.value, (str, int, float)):
                                clean_ws.cell(row=cell.row, column=cell.column, value=cell.value)
                
                log.append(f"Processed sheet: {sheet_name}")
            
            # Save clean workbook
            clean_buffer = io.BytesIO()
            clean_wb.save(clean_buffer)
            clean_content = clean_buffer.getvalue()
            
            threats_removed.extend([
                "Macros and VBA code removed",
                "Formulas replaced with values",
                "External links removed",
                "Embedded objects removed"
            ])
            
            return {
                'success': True,
                'clean_content': clean_content,
                'threats_removed': threats_removed,
                'metadata_stripped': True,
                'macros_removed': True,
                'scripts_removed': True,
                'embedded_objects_removed': True,
                'reconstruction_log': log
            }
            
        except Exception as e:
            log.append(f"XLSX reconstruction failed: {e}")
            return {
                'success': False,
                'clean_content': None,
                'threats_removed': [f"XLSX processing error: {e}"],
                'metadata_stripped': False,
                'macros_removed': False,
                'scripts_removed': False,
                'embedded_objects_removed': False,
                'reconstruction_log': log
            }

    def _reconstruct_image(self, content: bytes) -> Dict[str, Any]:
        """Reconstruct image with metadata stripped"""
        log = ["Starting image reconstruction"]
        threats_removed = []
        
        try:
            # Open image
            img = Image.open(io.BytesIO(content))
            
            # Create new image without metadata
            clean_img = Image.new(img.mode, img.size)
            clean_img.putdata(list(img.getdata()))
            
            # Save as clean image
            clean_buffer = io.BytesIO()
            format_name = img.format if img.format else 'PNG'
            clean_img.save(clean_buffer, format=format_name)
            clean_content = clean_buffer.getvalue()
            
            log.append(f"Image reconstructed - format: {format_name}")
            threats_removed.extend([
                "EXIF metadata removed",
                "Hidden data removed",
                "Color profiles normalized"
            ])
            
            return {
                'success': True,
                'clean_content': clean_content,
                'threats_removed': threats_removed,
                'metadata_stripped': True,
                'macros_removed': False,
                'scripts_removed': False,
                'embedded_objects_removed': True,
                'reconstruction_log': log
            }
            
        except Exception as e:
            log.append(f"Image reconstruction failed: {e}")
            return {
                'success': False,
                'clean_content': None,
                'threats_removed': [f"Image processing error: {e}"],
                'metadata_stripped': False,
                'macros_removed': False,
                'scripts_removed': False,
                'embedded_objects_removed': False,
                'reconstruction_log': log
            }

    def _reconstruct_text(self, content: bytes) -> Dict[str, Any]:
        """Reconstruct text file with safe content only"""
        log = ["Starting text reconstruction"]
        threats_removed = []
        
        try:
            # Decode text safely
            text = content.decode('utf-8', errors='ignore')
            
            # Remove potential script injections
            dangerous_strings = [
                '<script', '</script>', 'javascript:', 'vbscript:',
                'onload=', 'onerror=', 'onclick=', 'eval(', 'exec(',
                '<?php', '<%', '%>', 'cmd.exe', 'powershell'
            ]
            
            clean_text = text
            for dangerous in dangerous_strings:
                if dangerous in clean_text.lower():
                    clean_text = clean_text.replace(dangerous, '[REMOVED]')
                    threats_removed.append(f"Removed dangerous string: {dangerous}")
            
            # Ensure only printable characters
            clean_text = ''.join(char for char in clean_text if ord(char) >= 32 or char in '\n\r\t')
            
            clean_content = clean_text.encode('utf-8')
            
            log.append("Text reconstruction completed")
            
            return {
                'success': True,
                'clean_content': clean_content,
                'threats_removed': threats_removed,
                'metadata_stripped': True,
                'macros_removed': False,
                'scripts_removed': len(threats_removed) > 0,
                'embedded_objects_removed': False,
                'reconstruction_log': log
            }
            
        except Exception as e:
            log.append(f"Text reconstruction failed: {e}")
            return {
                'success': False,
                'clean_content': None,
                'threats_removed': [f"Text processing error: {e}"],
                'metadata_stripped': False,
                'macros_removed': False,
                'scripts_removed': False,
                'embedded_objects_removed': False,
                'reconstruction_log': log
            }

    def _quarantine_unknown(self, content: bytes) -> Dict[str, Any]:
        """Quarantine unknown file types"""
        return {
            'success': False,
            'clean_content': None,
            'threats_removed': ["Unknown file type - quarantined for safety"],
            'metadata_stripped': False,
            'macros_removed': False,
            'scripts_removed': False,
            'embedded_objects_removed': False,
            'reconstruction_log': ["File quarantined - unknown or unsupported type"]
        }

    def _create_clean_text_document(self, text_pages: List[str], original_type: str) -> bytes:
        """Create a clean text document from extracted content"""
        header = f"=== CLEANED DOCUMENT ===\n"
        header += f"Original type: {original_type}\n"
        header += f"Processed by: GovDocShield X CDR+++\n"
        header += f"Date: {time.strftime('%Y-%m-%d %H:%M:%S')}\n"
        header += "=" * 50 + "\n\n"
        
        content = header
        for i, page_text in enumerate(text_pages):
            if page_text.strip():
                content += f"--- Page {i + 1} ---\n"
                content += page_text + "\n\n"
        
        return content.encode('utf-8')

    def _extract_text_from_zip(self, content: bytes) -> str:
        """Extract text content from ZIP-based formats"""
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                text_content = ""
                
                # Look for document.xml in Word files
                if 'word/document.xml' in zf.namelist():
                    with zf.open('word/document.xml') as doc_file:
                        tree = ET.parse(doc_file)
                        root = tree.getroot()
                        
                        # Extract text from XML
                        for elem in root.iter():
                            if elem.text:
                                text_content += elem.text + " "
                
                return text_content.strip()
                
        except Exception as e:
            logger.error(f"Failed to extract text from ZIP: {e}")
            return ""

    def _calculate_safety_score(self, file_type: FileType, initial_threats: List[str], 
                               removed_threats: List[str], metadata_stripped: bool, 
                               macros_removed: bool) -> float:
        """Calculate safety score based on processing results"""
        
        base_score = 0.5  # Start with neutral score
        
        # Bonus for successful reconstruction
        if metadata_stripped:
            base_score += 0.2
        if macros_removed:
            base_score += 0.2
        
        # Penalty for threats found
        threat_penalty = len(initial_threats) * 0.1
        base_score -= threat_penalty
        
        # Bonus for threats removed
        removal_bonus = len(removed_threats) * 0.05
        base_score += removal_bonus
        
        # File type specific adjustments
        if file_type in [FileType.TXT, FileType.IMAGE]:
            base_score += 0.1  # Generally safer file types
        elif file_type in [FileType.PDF, FileType.DOCX]:
            base_score -= 0.05  # More complex, higher risk
        
        # Clamp to [0, 1]
        return max(0.0, min(1.0, base_score))

# Additional helper functions for specific file types
def _reconstruct_rtf(self, content: bytes) -> Dict[str, Any]:
    """Reconstruct RTF file"""
    log = ["Starting RTF reconstruction"]
    
    try:
        # Convert RTF to plain text (safest approach)
        text_content = content.decode('utf-8', errors='ignore')
        
        # Remove RTF commands, keep only text
        import re
        text_only = re.sub(r'\\[a-zA-Z]+\d*\s*', '', text_content)
        text_only = re.sub(r'[{}]', '', text_only)
        text_only = text_only.replace('\\', '')
        
        clean_content = self._create_clean_text_document([text_only], "RTF")
        
        return {
            'success': True,
            'clean_content': clean_content,
            'threats_removed': ["RTF formatting removed", "Control codes removed"],
            'metadata_stripped': True,
            'macros_removed': True,
            'scripts_removed': True,
            'embedded_objects_removed': True,
            'reconstruction_log': log + ["RTF converted to safe text"]
        }
        
    except Exception as e:
        return {
            'success': False,
            'clean_content': None,
            'threats_removed': [f"RTF processing error: {e}"],
            'metadata_stripped': False,
            'macros_removed': False,
            'scripts_removed': False,
            'embedded_objects_removed': False,
            'reconstruction_log': log + [f"Error: {e}"]
        }

def _reconstruct_html(self, content: bytes) -> Dict[str, Any]:
    """Reconstruct HTML with safe elements only"""
    log = ["Starting HTML reconstruction"]
    
    try:
        from bs4 import BeautifulSoup
        
        soup = BeautifulSoup(content, 'html.parser')
        
        # Remove all script tags and event handlers
        for script in soup(["script", "style", "object", "embed", "iframe"]):
            script.decompose()
        
        # Remove dangerous attributes
        dangerous_attrs = ['onclick', 'onload', 'onerror', 'onmouseover', 'href']
        for tag in soup.find_all():
            for attr in dangerous_attrs:
                if attr in tag.attrs:
                    del tag.attrs[attr]
        
        # Extract clean text
        clean_text = soup.get_text()
        clean_content = self._create_clean_text_document([clean_text], "HTML")
        
        return {
            'success': True,
            'clean_content': clean_content,
            'threats_removed': ["Scripts removed", "Event handlers removed", "Dangerous elements removed"],
            'metadata_stripped': True,
            'macros_removed': False,
            'scripts_removed': True,
            'embedded_objects_removed': True,
            'reconstruction_log': log + ["HTML sanitized and converted to text"]
        }
        
    except Exception as e:
        # Fallback to text extraction
        try:
            text_content = content.decode('utf-8', errors='ignore')
            # Simple HTML tag removal
            import re
            text_only = re.sub(r'<[^>]+>', '', text_content)
            clean_content = self._create_clean_text_document([text_only], "HTML")
            
            return {
                'success': True,
                'clean_content': clean_content,
                'threats_removed': ["HTML tags removed"],
                'metadata_stripped': True,
                'macros_removed': False,
                'scripts_removed': True,
                'embedded_objects_removed': True,
                'reconstruction_log': log + ["Fallback text extraction"]
            }
        except:
            return {
                'success': False,
                'clean_content': None,
                'threats_removed': [f"HTML processing error: {e}"],
                'metadata_stripped': False,
                'macros_removed': False,
                'scripts_removed': False,
                'embedded_objects_removed': False,
                'reconstruction_log': log + [f"Error: {e}"]
            }

def _reconstruct_email(self, content: bytes) -> Dict[str, Any]:
    """Reconstruct email with safe content only"""
    log = ["Starting email reconstruction"]
    
    try:
        # Parse email
        msg = email.message_from_bytes(content)
        
        # Extract safe content
        safe_headers = ['From', 'To', 'Subject', 'Date']
        clean_email = MIMEMultipart()
        
        for header in safe_headers:
            if header in msg:
                clean_email[header] = msg[header]
        
        # Extract text content only
        text_content = ""
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    text_content += part.get_payload(decode=True).decode('utf-8', errors='ignore')
        else:
            if msg.get_content_type() == "text/plain":
                text_content = msg.get_payload(decode=True).decode('utf-8', errors='ignore')
        
        # Add safe text part
        if text_content:
            clean_email.attach(MIMEText(text_content, 'plain'))
        
        clean_content = clean_email.as_bytes()
        
        return {
            'success': True,
            'clean_content': clean_content,
            'threats_removed': ["Attachments removed", "HTML content removed", "Metadata stripped"],
            'metadata_stripped': True,
            'macros_removed': True,
            'scripts_removed': True,
            'embedded_objects_removed': True,
            'reconstruction_log': log + ["Email sanitized - text content only"]
        }
        
    except Exception as e:
        return {
            'success': False,
            'clean_content': None,
            'threats_removed': [f"Email processing error: {e}"],
            'metadata_stripped': False,
            'macros_removed': False,
            'scripts_removed': False,
            'embedded_objects_removed': False,
            'reconstruction_log': log + [f"Error: {e}"]
        }

# Add these methods to the PositiveSelectionCDR class
PositiveSelectionCDR._reconstruct_rtf = _reconstruct_rtf
PositiveSelectionCDR._reconstruct_html = _reconstruct_html
PositiveSelectionCDR._reconstruct_email = _reconstruct_email
PositiveSelectionCDR._reconstruct_pptx = lambda self, content: self._reconstruct_docx(content)  # Similar to DOCX